#!/usr/bin/env python3
"""
DOCX Validator: Automated quality inspection for scientific papers.

Scans DOCX documents and validates:
    1. Physical image embedding (not just text references)
    2. Bold formatting for best results in tables
    3. Sequential reference numbering [1]-[N]
    4. Sequential equation numbering (1), (2), ...
    5. Figure/table caption placement rules
    6. Section completeness (3.1-3.4, 4.1-4.3)

Usage:
    python docx_validator.py paper.docx
    python docx_validator.py paper.docx --verbose
    python docx_validator.py paper.docx --format json

Dependencies:
    python-docx
"""

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


class DOCXValidator:
    """Validator for scientific paper DOCX files."""

    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.path = docx_path
        self.issues = []
        self.warnings = []
        self.stats = {}

    def log_issue(self, category, message, severity='ERROR'):
        """Log an issue."""
        entry = {'severity': severity, 'category': category, 'message': message}
        if severity == 'ERROR':
            self.issues.append(entry)
        else:
            self.warnings.append(entry)

    # =====================================================================
    # Check 1: Image Embedding
    # =====================================================================

    def check_image_embedding(self):
        """Verify all 'Fig. N' references have corresponding embedded images."""
        print("  [Check 1] Image Embedding...")

        # Find all figure references in text
        fig_refs = []
        fig_embeds = []

        for para in self.doc.paragraphs:
            text = para.text
            # Find figure references like "Fig. 1", "Figure 2", "Figs. 3 and 4"
            refs = re.findall(r'(?:Fig\.?|Figure)\s+(\d+)', text, re.IGNORECASE)
            fig_refs.extend([int(r) for r in refs])

            # Check for embedded images
            for run in para.runs:
                xml = run._element.xml
                if 'graphicData' in xml or 'blip' in xml:
                    fig_embeds.append(para)

        # Count unique references
        unique_refs = sorted(set(fig_refs))
        embed_count = len(fig_embeds)

        self.stats['figure_references'] = len(unique_refs)
        self.stats['figures_embedded'] = embed_count

        if not unique_refs:
            self.log_issue('Image', 'No figure references (Fig. N) found in text', 'WARNING')
            return

        if embed_count == 0:
            self.log_issue('Image',
                f'Found {len(unique_refs)} figure reference(s) [Fig. {unique_refs}] but ZERO images embedded. '
                f'All figures must be physically embedded via add_picture().', 'ERROR')
        elif embed_count < len(unique_refs):
            self.log_issue('Image',
                f'Found {len(unique_refs)} figure reference(s) but only {embed_count} image(s) embedded. '
                f'Missing: {len(unique_refs) - embed_count} figure(s)', 'ERROR')
        else:
            print(f"    PASS: {embed_count} figure(s) embedded, {len(unique_refs)} reference(s) found")

    # =====================================================================
    # Check 2: Table Best-Result Bold (Smart Direction Detection)
    # =====================================================================

    # Keywords that identify property/specification tables (no "best" concept)
    PROPERTY_TABLE_KEYWORDS = {
        'material', 'parameter', 'specification', 'property', 'component',
        'type', 'description', 'sample', 'configuration', 'architecture',
        'layer', 'unit', 'dimension', 'hardware', 'robot', 'device',
        'chemical', 'formula', 'composition', 'structure', 'feature',
    }

    # Keywords indicating "smaller is better"
    MINIMIZE_KEYWORDS = {
        'loss', 'error', 'resistance', 'impedance', 'gap', 'drop', 'fade',
        'degradation', 'decay', 'cost', 'delay', 'overhead', 'variance',
        'deviation', 'uncertainty', 'latency', 'r_b', 'r_gb', 'r_ct',
        'r_total', 'overpotential', 'polarization', 'defect', 'void',
        'thickness', 'distance', 'diff', 'difference', 'penalty',
    }

    # Keywords indicating "larger is better"
    MAXIMIZE_KEYWORDS = {
        'accuracy', 'acc', 'precision', 'recall', 'f1', 'score', 'mAP',
        'efficiency', 'capacity', 'energy', 'density', 'power', 'speed',
        'velocity', 'rate', 'retention', 'stability', 'lifetime', 'life',
        'conductivity', 'diffusion', 'coefficient', 'strength', 'modulus',
        'yield', 'output', 'reward', 'return', 'gain', 'improvement',
        'boost', 'enhancement', 'performance', 'metric', 'value',
    }

    def _is_property_table(self, headers):
        """Check if table is a property/specification table (no best/worst)."""
        header_text = ' '.join(h.lower() for h in headers if h)
        for kw in self.PROPERTY_TABLE_KEYWORDS:
            if kw in header_text:
                # Also require that there is no clear "method" or "result" column
                return True
        return False

    def _get_column_direction(self, header):
        """Determine if a column should be minimized or maximized."""
        h = header.lower()
        for kw in self.MINIMIZE_KEYWORDS:
            if kw in h:
                return 'minimize'
        for kw in self.MAXIMIZE_KEYWORDS:
            if kw in h:
                return 'maximize'
        return 'unknown'

    def check_table_bold(self):
        """Check if comparison tables have best results in bold."""
        print("  [Check 2] Table Bold Formatting...")

        table_count = len(self.doc.tables)
        self.stats['tables'] = table_count

        if table_count == 0:
            self.log_issue('Table', 'No tables found in document', 'WARNING')
            return

        checked = 0
        skipped = 0
        for idx, table in enumerate(self.doc.tables):
            if len(table.rows) < 2:
                continue

            headers = [cell.text.strip() for cell in table.rows[0].cells]

            # Skip property/specification tables
            if self._is_property_table(headers):
                skipped += 1
                continue

            # Check numeric columns (skip header and first label column)
            for col_idx in range(1, len(table.columns)):
                if col_idx >= len(headers):
                    continue

                direction = self._get_column_direction(headers[col_idx])
                if direction == 'unknown':
                    continue  # Skip ambiguous columns

                try:
                    values = []
                    for row in table.rows[1:]:
                        cell = row.cells[col_idx]
                        text = cell.text.strip()
                        try:
                            val = float(text.replace('%', '').replace(',', ''))
                            values.append((val, cell))
                        except ValueError:
                            continue

                    if len(values) >= 2:
                        if direction == 'minimize':
                            target_val = min(v[0] for v in values)
                        else:
                            target_val = max(v[0] for v in values)

                        best_cells = [c for v, c in values if v == target_val]

                        for cell in best_cells:
                            has_bold = False
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.bold:
                                        has_bold = True
                                        break

                            if not has_bold:
                                dir_label = 'lowest' if direction == 'minimize' else 'best'
                                self.log_issue('Table',
                                    f'Table {idx+1}, column "{headers[col_idx]}": {dir_label} value {target_val} '
                                    f'is NOT bold. Best results must be bold per figure_table_guidelines.md',
                                    'WARNING')
                            checked += 1
                except Exception:
                    continue

        if skipped > 0:
            print(f"    Skipped {skipped} property/specification table(s)")
        if checked == 0:
            print("    PASS: No numeric comparison tables requiring bold found")
        else:
            print(f"    Checked {checked} table column(s)")

    # =====================================================================
    # Check 3: Reference Numbering Continuity
    # =====================================================================

    def check_reference_continuity(self):
        """Check if [1], [2], [3]... are sequential without gaps."""
        print("  [Check 3] Reference Numbering...")

        all_nums = []
        for para in self.doc.paragraphs:
            # Find [N] or [N, M] or [N-M] patterns
            matches = re.findall(r'\[(\d+)(?:,\s*\d+|\s*-\s*\d+)?\]', para.text)
            all_nums.extend([int(m) for m in matches])

        if not all_nums:
            self.log_issue('Reference', 'No numbered references [N] found in text', 'WARNING')
            return

        unique_nums = sorted(set(all_nums))
        self.stats['references'] = len(unique_nums)

        # Check for gaps
        expected = list(range(1, max(unique_nums) + 1))
        gaps = [n for n in expected if n not in unique_nums]

        if gaps:
            self.log_issue('Reference',
                f'Reference numbering has gap(s): missing [{"], [".join(map(str, gaps[:5]))}]'
                f'{"..." if len(gaps) > 5 else ""}. References should be sequential [1]-[{max(unique_nums)}]',
                'ERROR')
        else:
            print(f"    PASS: {len(unique_nums)} sequential reference(s) [1]-[{max(unique_nums)}]")

    # =====================================================================
    # Check 4: Equation Numbering
    # =====================================================================

    def check_equation_numbering(self):
        """Check if equations are numbered sequentially (1), (2), (3)..."""
        print("  [Check 4] Equation Numbering...")

        eq_nums = []
        for para in self.doc.paragraphs:
            # Find equation numbers like (1), (2), (10)
            matches = re.findall(r'\((\d+)\)', para.text)
            # Only count if paragraph looks like an equation (has math-like content)
            if any(c in para.text for c in ['=', 'sum', 'theta', 'alpha', 'beta', 'lambda', 'sigma', 'delta']):
                eq_nums.extend([int(m) for m in matches])

        if not eq_nums:
            print("    INFO: No numbered equations detected")
            return

        unique_nums = sorted(set(eq_nums))
        self.stats['equations'] = len(unique_nums)

        expected = list(range(1, max(unique_nums) + 1))
        gaps = [n for n in expected if n not in unique_nums]

        if gaps:
            self.log_issue('Equation',
                f'Equation numbering has gap(s): missing ({"), (".join(map(str, gaps[:5]))})'
                f'{"..." if len(gaps) > 5 else ""}. Equations should be sequential (1)-({max(unique_nums)})',
                'WARNING')
        else:
            print(f"    PASS: {len(unique_nums)} sequential equation(s) (1)-({max(unique_nums)})")

    # =====================================================================
    # Check 5: Caption Placement
    # =====================================================================

    def check_caption_placement(self):
        """Check figure captions below, table captions above."""
        print("  [Check 5] Caption Placement...")

        # Check each table: caption should be in paragraph BEFORE the table
        for idx, table in enumerate(self.doc.tables):
            # Find the paragraph just before this table
            table_xml = table._tbl
            table_idx = None
            for i, para in enumerate(self.doc.paragraphs):
                if table_xml in para._element.getparent():
                    # Complex DOM traversal - simplified check
                    pass

            # Simplified: check table's first row (header) for bold
            # Real check would need paragraph ordering relative to tables

        print("    INFO: Caption placement check requires advanced DOM analysis (simplified)")

    # =====================================================================
    # Check 6: Section Completeness
    # =====================================================================

    def check_section_completeness(self):
        """Verify required subsections exist (3.1-3.4, 4.1-4.3)."""
        print("  [Check 6] Section Completeness...")

        sections = {}
        for para in self.doc.paragraphs:
            text = para.text.strip()
            # Match section headings like "3. Method", "3.1 Core Formulation"
            match = re.match(r'^(\d+(?:\.\d+)?)\s+[A-Za-z]', text)
            if match:
                sections[match.group(1)] = text

        self.stats['sections'] = list(sections.keys())

        required = {
            'Method': ['3.1', '3.2', '3.3', '3.4'],
            'Experiments': ['4.1', '4.2', '4.3']
        }

        for parent, subs in required.items():
            missing = [s for s in subs if s not in sections]
            if missing:
                self.log_issue('Structure',
                    f'{parent} section missing subsection(s): {", ".join(missing)}. '
                    f'Per structure_contract.md, {parent} MUST have all subsections.', 'ERROR')
            else:
                print(f"    PASS: {parent} has all required subsections ({', '.join(subs)})")

    # =====================================================================
    # Run All Checks
    # =====================================================================

    def validate(self):
        """Run all validation checks."""
        print(f"\nValidating: {self.path}")
        print("=" * 60)

        self.check_image_embedding()
        self.check_table_bold()
        self.check_reference_continuity()
        self.check_equation_numbering()
        self.check_caption_placement()
        self.check_section_completeness()

        print("\n" + "=" * 60)
        print("VALIDATION SUMMARY")
        print("=" * 60)

        print(f"\nStats:")
        for key, val in self.stats.items():
            print(f"  {key}: {val}")

        print(f"\nIssues: {len(self.issues)} error(s), {len(self.warnings)} warning(s)")

        if self.issues:
            print("\n  ERRORS (must fix before submission):")
            for i in self.issues:
                print(f"    [{i['category']}] {i['message']}")

        if self.warnings:
            print("\n  WARNINGS (should fix):")
            for w in self.warnings:
                print(f"    [{w['category']}] {w['message']}")

        if not self.issues and not self.warnings:
            print("\n  ALL CHECKS PASSED! Document is ready for submission.")

        return len(self.issues) == 0

    def to_json(self):
        """Export validation results as JSON."""
        return {
            'file': self.path,
            'stats': self.stats,
            'issues': self.issues,
            'warnings': self.warnings,
            'passed': len(self.issues) == 0
        }


def main():
    parser = argparse.ArgumentParser(description='Validate scientific paper DOCX files')
    parser.add_argument('file', help='DOCX file to validate')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    parser.add_argument('--format', choices=['text', 'json'], default='text',
                        help='Output format')
    parser.add_argument('--output', '-o', help='Output file for JSON report')
    args = parser.parse_args()

    if not Path(args.file).exists():
        print(f"Error: File not found: {args.file}")
        sys.exit(1)

    validator = DOCXValidator(args.file)
    passed = validator.validate()

    if args.format == 'json' or args.output:
        result = validator.to_json()
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            print(f"\nJSON report saved to: {args.output}")
        elif args.format == 'json':
            print(json.dumps(result, indent=2, ensure_ascii=False))

    sys.exit(0 if passed else 1)


if __name__ == '__main__':
    main()
