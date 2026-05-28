#!/usr/bin/env python3
"""
DOCX to LaTeX Converter: Convert scientific papers to camera-ready LaTeX.

Supports venue-specific templates (CVPR, NeurIPS, ICML, ICLR, ACL).
When pandoc is unavailable, falls back to native python-docx parsing.

Usage:
    python docx_to_latex.py paper.docx --venue cvpr
    python docx_to_latex.py paper.docx --venue neurips --output paper.tex
    python docx_to_latex.py paper.docx --template cvpr2024.sty

Dependencies:
    python-docx
"""

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# Venue Templates
# =============================================================================

VENUE_TEMPLATES = {
    'cvpr': {
        'documentclass': r'\documentclass[10pt,twocolumn,letterpaper]{article}',
        'packages': r'''
\usepackage{cvpr}
\usepackage{times}
\usepackage{epsfig}
\usepackage{graphicx}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{booktabs}
\usepackage{color}''',
        'columns': 2,
        'page': 'letter',
    },
    'neurips': {
        'documentclass': r'\documentclass{article}',
        'packages': r'''
\usepackage[preprint]{neurips_2024}
\usepackage{times}
\usepackage{hyperref}
\usepackage{url}
\usepackage{booktabs}
\usepackage{amsfonts}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{graphicx}
\usepackage{color}''',
        'columns': 1,
        'page': 'letter',
    },
    'icml': {
        'documentclass': r'\documentclass{article}',
        'packages': r'''
\usepackage[accepted]{icml2024}
\usepackage{times}
\usepackage{hyperref}
\usepackage{url}
\usepackage{booktabs}
\usepackage{amsfonts}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{graphicx}
\usepackage{color}''',
        'columns': 1,
        'page': 'letter',
    },
    'iclr': {
        'documentclass': r'\documentclass{article}',
        'packages': r'''
\usepackage[accepted]{iclr2024}
\usepackage{times}
\usepackage{hyperref}
\usepackage{url}
\usepackage{booktabs}
\usepackage{amsfonts}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{graphicx}
\usepackage{color}''',
        'columns': 2,
        'page': 'letter',
    },
    'acl': {
        'documentclass': r'\documentclass[11pt]{article}',
        'packages': r'''
\usepackage[review]{acl}
\usepackage{times}
\usepackage{latexsym}
\usepackage{hyperref}
\usepackage{url}
\usepackage{booktabs}
\usepackage{graphicx}
\usepackage{color}''',
        'columns': 2,
        'page': 'letter',
    },
}


def has_pandoc():
    """Check if pandoc is available."""
    try:
        subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def convert_with_pandoc(docx_path, output_path, venue='cvpr'):
    """Convert DOCX to LaTeX using pandoc."""
    print(f"Using pandoc for conversion...")

    # First convert to raw LaTeX
    temp_tex = output_path.replace('.tex', '_raw.tex')

    cmd = [
        'pandoc', docx_path,
        '-f', 'docx',
        '-t', 'latex',
        '--wrap=none',
        '-o', temp_tex
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Pandoc error: {result.stderr}")
        return False

    # Post-process: wrap in venue template
    with open(temp_tex, 'r', encoding='utf-8') as f:
        content = f.read()

    # Clean up pandoc output
    content = _post_process_latex(content, venue)

    # Wrap in template
    template = VENUE_TEMPLATES.get(venue, VENUE_TEMPLATES['cvpr'])
    full_tex = f"""{template['documentclass']}
{template['packages']}

{content}

\end{{document}}
"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_tex)

    os.remove(temp_tex)
    print(f"LaTeX saved to: {output_path}")
    return True


def _post_process_latex(content, venue):
    """Post-process pandoc LaTeX output."""
    # Fix figure placement
    content = re.sub(r'\\begin\{figure\}', r'\\begin{figure}[t]', content)
    content = re.sub(r'\\begin\{table\}', r'\\begin{table}[t]', content)

    # Fix citation format [N] -> \cite{N}
    content = re.sub(r'\[(\d+)\]', r'\\cite{\1}', content)

    # Fix equation numbering
    content = re.sub(r'\((\d+)\)\s*\\end\{equation\}',
                     r'\\label{eq:\1}\\end{equation}', content)

    return content


# =============================================================================
# Native DOCX → LaTeX Parser
# =============================================================================

def convert_native(docx_path, output_path, venue='cvpr'):
    """Convert DOCX to LaTeX using native python-docx parsing (no pandoc)."""
    print(f"Using native converter (pandoc not available)...")

    doc = Document(docx_path)
    template = VENUE_TEMPLATES.get(venue, VENUE_TEMPLATES['cvpr'])

    lines = []
    lines.append(template['documentclass'])
    lines.append(template['packages'])
    lines.append('')
    lines.append(r'\begin{document}')
    lines.append('')

    # Process paragraphs
    for para in doc.paragraphs:
        latex = _paragraph_to_latex(para)
        if latex:
            lines.append(latex)

    # Process tables
    for table in doc.tables:
        latex = _table_to_latex(table)
        if latex:
            lines.append(latex)

    lines.append('')
    lines.append(r'\end{document}')

    content = '\n'.join(lines)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

    print(f"LaTeX saved to: {output_path}")
    return True


def _paragraph_to_latex(para):
    """Convert a single paragraph to LaTeX."""
    text = para.text.strip()
    if not text:
        return ''

    # Detect section headings
    heading_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
    if heading_match:
        level = heading_match.group(1)
        title = heading_match.group(2)

        if '.' not in level:  # Top-level section
            return f"\\section{{{title}}}"
        elif level.count('.') == 1:  # Subsection
            return f"\\subsection{{{title}}}"
        else:
            return f"\\subsubsection{{{title}}}"

    # Detect title (first paragraph, large font)
    if para.runs and para.runs[0].bold and para.runs[0].font.size and para.runs[0].font.size.pt >= 16:
        return f"\\title{{{text}}}\n\\maketitle"

    # Detect abstract heading
    if text.lower() == 'abstract':
        return r'\begin{abstract}'

    # Detect figure caption
    if text.startswith('Figure '):
        return f"\\caption{{{text[8:]}}}"  # Remove "Figure N. "

    # Detect table caption
    if text.startswith('Table '):
        return f"\\caption{{{text[8:]}}}"

    # Detect equation
    eq_match = re.search(r'(.+?)\s*\((\d+)\)\s*$', text)
    if eq_match and '=' in text:
        eq_body = eq_match.group(1).strip()
        eq_num = eq_match.group(2)
        return f"\\begin{{equation}}\n{eq_body} \n\\label{{eq:{eq_num}}}\n\\end{{equation}}"

    # Detect references section
    if text.startswith('[') and re.match(r'\[\d+\]', text):
        # Reference entry
        return f"\\bibitem{{{text[1:text.index(']')]}}} {text[text.index(']')+1:].strip()}"

    # Regular paragraph
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return f"\\begin{{center}}\n{text}\n\\end{{center}}"
    else:
        return text


def _table_to_latex(table):
    """Convert a table to LaTeX."""
    n_cols = len(table.columns)
    col_spec = 'l' + 'c' * (n_cols - 1)

    lines = [r'\begin{table}[t]', r'\centering']

    lines.append(r'\begin{tabular}{' + col_spec + '}')
    lines.append(r'\toprule')

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        # Bold first row (header)
        if i == 0:
            cells = [f'\\textbf{{{c}}}' for c in cells]
        line = ' & '.join(cells) + r' \\ '
        lines.append(line)
        if i == 0:
            lines.append(r'\midrule')

    lines.append(r'\bottomrule')
    lines.append(r'\end{tabular}')
    lines.append(r'\end{table}')

    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX to LaTeX for camera-ready submission')
    parser.add_argument('file', help='Input DOCX file')
    parser.add_argument('--venue', choices=list(VENUE_TEMPLATES.keys()), default='cvpr',
                        help='Target conference venue')
    parser.add_argument('--template', help='Custom LaTeX template file (overrides --venue)')
    parser.add_argument('--output', '-o', help='Output .tex file')
    parser.add_argument('--force-native', action='store_true',
                        help='Force native conversion even if pandoc is available')
    args = parser.parse_args()

    if not Path(args.file).exists():
        print(f"Error: File not found: {args.file}")
        sys.exit(1)

    output_path = args.output or args.file.replace('.docx', f'_{args.venue}.tex')

    print(f"Converting: {args.file}")
    print(f"Venue: {args.venue.upper()}")
    print(f"Output: {output_path}")
    print("-" * 60)

    if not args.force_native and has_pandoc():
        success = convert_with_pandoc(args.file, output_path, args.venue)
    else:
        success = convert_native(args.file, output_path, args.venue)

    if success:
        print("\nConversion completed successfully!")
        print(f"Next steps:")
        print(f"  1. Download the official {args.venue.upper()} LaTeX template")
        print(f"  2. Place {output_path} in the template directory")
        print(f"  3. Run: pdflatex {Path(output_path).name}")
    else:
        print("\nConversion failed.")
        sys.exit(1)


if __name__ == '__main__':
    main()
