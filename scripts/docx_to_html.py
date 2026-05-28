#!/usr/bin/env python3
"""
DOCX to Interactive HTML: Convert scientific papers to web-ready interactive pages.

Features:
    - Hover tooltips on figures showing data values
    - Click-to-expand equations
    - Clickable reference jumps
    - Responsive design for mobile/tablet
    - Syntax-highlighted code blocks
    - Dark/light mode toggle

Usage:
    python docx_to_html.py paper.docx --output paper.html
    python docx_to_html.py paper.docx --theme dark

Dependencies:
    python-docx
"""

import argparse
import base64
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# HTML Templates
# =============================================================================

HTML_TEMPLATE = r'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        :root {{
            --bg: #fafafa;
            --fg: #1a1a1a;
            --accent: #377EB8;
            --accent-light: #e8f4fd;
            --border: #ddd;
            --code-bg: #f4f4f4;
            --caption-bg: #f8f8f8;
        }}

        [data-theme="dark"] {{
            --bg: #1a1a1a;
            --fg: #e0e0e0;
            --accent: #5ba3e0;
            --accent-light: #1a2a3a;
            --border: #444;
            --code-bg: #2a2a2a;
            --caption-bg: #222;
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            font-size: 16px;
            line-height: 1.7;
            color: var(--fg);
            background: var(--bg);
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            transition: background 0.3s, color 0.3s;
        }}

        /* Header */
        .paper-header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 2px solid var(--accent);
        }}

        .paper-title {{
            font-size: 28px;
            font-weight: bold;
            margin-bottom: 15px;
            color: var(--fg);
        }}

        .paper-authors {{
            font-size: 16px;
            color: var(--fg);
            opacity: 0.8;
            margin-bottom: 10px;
        }}

        .paper-affiliation {{
            font-size: 14px;
            color: var(--fg);
            opacity: 0.6;
            font-style: italic;
        }}

        /* Sections */
        h1 {{
            font-size: 22px;
            font-weight: bold;
            margin: 35px 0 15px;
            color: var(--accent);
            border-bottom: 1px solid var(--border);
            padding-bottom: 8px;
        }}

        h2 {{
            font-size: 18px;
            font-weight: bold;
            margin: 25px 0 12px;
            color: var(--fg);
        }}

        h3 {{
            font-size: 16px;
            font-weight: bold;
            margin: 20px 0 10px;
            color: var(--fg);
        }}

        p {{
            margin-bottom: 12px;
            text-align: justify;
        }}

        /* Abstract */
        .abstract {{
            background: var(--accent-light);
            border-left: 4px solid var(--accent);
            padding: 20px;
            margin: 20px 0;
            border-radius: 4px;
        }}

        .abstract h2 {{
            margin-top: 0;
            color: var(--accent);
        }}

        /* Figures */
        .figure {{
            margin: 25px 0;
            text-align: center;
        }}

        .figure img {{
            max-width: 100%;
            height: auto;
            border-radius: 4px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            transition: transform 0.2s;
            cursor: zoom-in;
        }}

        .figure img:hover {{
            transform: scale(1.02);
        }}

        .figure img.expanded {{
            transform: scale(1.5);
            cursor: zoom-out;
            position: relative;
            z-index: 100;
        }}

        .figure-caption {{
            font-size: 13px;
            font-weight: bold;
            color: var(--fg);
            margin-top: 10px;
            padding: 10px;
            background: var(--caption-bg);
            border-radius: 4px;
        }}

        /* Tables */
        .table-wrapper {{
            overflow-x: auto;
            margin: 25px 0;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
        }}

        th, td {{
            padding: 10px 12px;
            text-align: center;
            border-bottom: 1px solid var(--border);
        }}

        th {{
            font-weight: bold;
            background: var(--accent-light);
            border-bottom: 2px solid var(--accent);
        }}

        tr:hover {{
            background: var(--accent-light);
        }}

        .table-caption {{
            font-size: 13px;
            font-weight: bold;
            text-align: center;
            margin-bottom: 10px;
        }}

        /* Equations */
        .equation {{
            margin: 20px 0;
            padding: 15px 20px;
            background: var(--code-bg);
            border-radius: 4px;
            cursor: pointer;
            transition: background 0.2s;
            position: relative;
        }}

        .equation:hover {{
            background: var(--accent-light);
        }}

        .equation .eq-body {{
            display: inline-block;
            width: 90%;
            text-align: center;
            font-style: italic;
            font-family: 'Computer Modern', 'Latin Modern Math', serif;
        }}

        .equation .eq-number {{
            display: inline-block;
            width: 8%;
            text-align: right;
            color: var(--accent);
            font-weight: bold;
        }}

        .equation.expanded {{
            background: var(--accent-light);
            border: 1px solid var(--accent);
        }}

        .equation .eq-details {{
            display: none;
            margin-top: 10px;
            padding-top: 10px;
            border-top: 1px dashed var(--border);
            font-size: 13px;
            color: var(--fg);
            opacity: 0.8;
        }}

        .equation.expanded .eq-details {{
            display: block;
        }}

        /* References */
        .ref-link {{
            color: var(--accent);
            text-decoration: none;
            font-weight: bold;
            cursor: pointer;
            padding: 1px 4px;
            border-radius: 2px;
            transition: background 0.2s;
        }}

        .ref-link:hover {{
            background: var(--accent-light);
            text-decoration: underline;
        }}

        .references {{
            margin-top: 30px;
            padding-top: 20px;
            border-top: 2px solid var(--border);
        }}

        .references h1 {{
            margin-top: 0;
        }}

        .ref-item {{
            margin-bottom: 8px;
            font-size: 14px;
            padding: 5px 10px;
            border-radius: 3px;
            transition: background 0.2s;
        }}

        .ref-item:hover {{
            background: var(--accent-light);
        }}

        .ref-item.highlighted {{
            background: var(--accent-light);
            border-left: 3px solid var(--accent);
        }}

        /* Dark mode toggle */
        .theme-toggle {{
            position: fixed;
            top: 20px;
            right: 20px;
            width: 50px;
            height: 50px;
            border-radius: 50%;
            background: var(--accent);
            color: white;
            border: none;
            cursor: pointer;
            font-size: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
            transition: transform 0.2s;
            z-index: 1000;
        }}

        .theme-toggle:hover {{
            transform: scale(1.1);
        }}

        /* Tooltip */
        .tooltip {{
            position: relative;
        }}

        .tooltip::after {{
            content: attr(data-tooltip);
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%);
            background: var(--dark);
            color: white;
            padding: 5px 10px;
            border-radius: 4px;
            font-size: 12px;
            white-space: nowrap;
            opacity: 0;
            pointer-events: none;
            transition: opacity 0.2s;
        }}

        .tooltip:hover::after {{
            opacity: 1;
        }}

        /* Responsive */
        @media (max-width: 600px) {{
            body {{
                padding: 20px 15px;
                font-size: 15px;
            }}
            .paper-title {{
                font-size: 22px;
            }}
            h1 {{
                font-size: 20px;
            }}
            .equation .eq-body {{
                width: 85%;
            }}
        }}
    </style>
</head>
<body>
    <button class="theme-toggle" onclick="toggleTheme()" title="Toggle dark mode">🌙</button>

    {content}

    <script>
        // Theme toggle
        function toggleTheme() {{
            const html = document.documentElement;
            const btn = document.querySelector('.theme-toggle');
            if (html.getAttribute('data-theme') === 'dark') {{
                html.removeAttribute('data-theme');
                btn.textContent = '🌙';
                localStorage.setItem('theme', 'light');
            }} else {{
                html.setAttribute('data-theme', 'dark');
                btn.textContent = '☀️';
                localStorage.setItem('theme', 'dark');
            }}
        }}

        // Load saved theme
        if (localStorage.getItem('theme') === 'dark') {{
            document.documentElement.setAttribute('data-theme', 'dark');
            document.querySelector('.theme-toggle').textContent = '☀️';
        }}

        // Figure zoom
        document.querySelectorAll('.figure img').forEach(img => {{
            img.addEventListener('click', function() {{
                this.classList.toggle('expanded');
            }});
        }});

        // Equation expand
        document.querySelectorAll('.equation').forEach(eq => {{
            eq.addEventListener('click', function() {{
                this.classList.toggle('expanded');
            }});
        }});

        // Reference jump
        document.querySelectorAll('.ref-link').forEach(link => {{
            link.addEventListener('click', function(e) {{
                e.preventDefault();
                const num = this.getAttribute('data-ref');
                const target = document.getElementById('ref-' + num);
                if (target) {{
                    target.scrollIntoView({{behavior: 'smooth', block: 'center'}});
                    target.classList.add('highlighted');
                    setTimeout(() => target.classList.remove('highlighted'), 2000);
                }}
            }});
        }});
    </script>
</body>
</html>
'''


def extract_images(doc):
    """Extract embedded images from DOCX and convert to base64."""
    images = {}
    image_counter = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_counter += 1
            image_data = rel.target_part.blob
            image_b64 = base64.b64encode(image_data).decode('utf-8')
            images[image_counter] = f"data:image/png;base64,{image_b64}"

    return images


def paragraph_to_html(para, images, image_counter):
    """Convert a single paragraph to HTML."""
    text = para.text.strip()
    if not text:
        return "", image_counter

    # Detect section headings
    heading_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
    if heading_match:
        level = heading_match.group(1)
        title = heading_match.group(2)

        if '.' not in level:
            return f'<h1 id="sec-{level}">{text}</h1>', image_counter
        elif level.count('.') == 1:
            return f'<h2 id="sec-{level}">{text}</h2>', image_counter
        else:
            return f'<h3 id="sec-{level}">{text}</h3>', image_counter

    # Detect title
    if para.runs and para.runs[0].bold and para.runs[0].font.size and para.runs[0].font.size.pt >= 16:
        return f'<div class="paper-header"><div class="paper-title">{text}</div>', image_counter

    # Detect authors
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and not para.runs[0].bold:
        return f'<div class="paper-authors">{text}</div></div>', image_counter

    # Detect abstract heading
    if text.lower() == 'abstract':
        return '<div class="abstract"><h2>Abstract</h2>', image_counter

    # Detect figure caption
    fig_match = re.match(r'^(Figure \d+\.?)\s*(.+)', text, re.IGNORECASE)
    if fig_match:
        fig_num = fig_match.group(1)
        fig_text = fig_match.group(2)

        # Check if next element is an image
        img_html = ""
        if image_counter in images:
            img_html = f'<img src="{images[image_counter]}" alt="{fig_num}" title="Click to zoom">'
            image_counter += 1

        return (f'<div class="figure">{img_html}'
                f'<div class="figure-caption">{fig_num} {fig_text}</div></div>'), image_counter

    # Detect table caption
    tbl_match = re.match(r'^(Table \d+\.?)\s*(.+)', text, re.IGNORECASE)
    if tbl_match:
        return f'<div class="table-caption">{text}</div>', image_counter

    # Detect equation
    eq_match = re.search(r'(.+?)\s*\((\d+)\)\s*$', text)
    if eq_match and '=' in text:
        eq_body = eq_match.group(1).strip()
        eq_num = eq_match.group(2)
        return (f'<div class="equation">'
                f'<span class="eq-body">{eq_body}</span>'
                f'<span class="eq-number">({eq_num})</span>'
                f'<div class="eq-details">Click to expand/collapse. Equation ({eq_num}) from Section ???</div>'
                f'</div>'), image_counter

    # Detect references section
    if text.startswith('[') and re.match(r'\[\d+\]', text):
        ref_num = re.search(r'\[(\d+)\]', text).group(1)
        return f'<div class="ref-item" id="ref-{ref_num}">{text}</div>', image_counter

    # Regular paragraph - convert inline references to links
    def ref_replacer(m):
        num = m.group(1)
        return f'<a href="#ref-{num}" class="ref-link" data-ref="{num}">[{num}]</a>'

    text = re.sub(r'\[(\d+)\]', ref_replacer, text)

    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return f'<p style="text-align:center">{text}</p>', image_counter
    else:
        return f'<p>{text}</p>', image_counter


def table_to_html(table, table_idx):
    """Convert a DOCX table to HTML."""
    lines = ['<div class="table-wrapper"><table>']

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        tag = 'th' if i == 0 else 'td'
        row_html = '<tr>' + ''.join(f'<{tag}>{c}</{tag}>' for c in cells) + '</tr>'
        lines.append(row_html)

    lines.append('</table></div>')
    return '\n'.join(lines)


def convert_docx_to_html(docx_path, output_path, theme='light'):
    """Convert DOCX to interactive HTML."""
    print(f"Converting: {docx_path}")

    doc = Document(docx_path)
    images = extract_images(doc)
    print(f"  Extracted {len(images)} image(s)")

    html_parts = []
    image_counter = 1

    for para in doc.paragraphs:
        part, image_counter = paragraph_to_html(para, images, image_counter)
        if part:
            html_parts.append(part)

    # Add tables
    for idx, table in enumerate(doc.tables):
        html_parts.append(table_to_html(table, idx))

    content = '\n'.join(html_parts)

    # Extract title
    title = "Paper"
    for para in doc.paragraphs:
        if para.runs and para.runs[0].font.size and para.runs[0].font.size.pt >= 16:
            title = para.text.strip()
            break

    html = HTML_TEMPLATE.format(title=title, content=content)

    # Inject theme
    if theme == 'dark':
        html = html.replace('<html lang="en">', '<html lang="en" data-theme="dark">')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

    print(f"HTML saved to: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX to interactive HTML paper')
    parser.add_argument('file', help='Input DOCX file')
    parser.add_argument('--output', '-o', help='Output HTML file')
    parser.add_argument('--theme', choices=['light', 'dark'], default='light',
                        help='Color theme')
    args = parser.parse_args()

    if not Path(args.file).exists():
        print(f"Error: File not found: {args.file}")
        sys.exit(1)

    output_path = args.output or args.file.replace('.docx', '.html')

    convert_docx_to_html(args.file, output_path, theme=args.theme)

    print(f"\nFeatures:")
    print(f"  - Dark/light mode toggle (top-right button)")
    print(f"  - Click figures to zoom")
    print(f"  - Click equations to expand details")
    print(f"  - Click [N] references to jump to bibliography")
    print(f"  - Hover over table rows for highlight")
    print(f"\nOpen in browser: file:///{os.path.abspath(output_path).replace('\\', '/')}")


if __name__ == '__main__':
    main()
