# SCI Paper Writing — Reusable Code Templates
#
# These templates provide production-ready code skeletons for generating
# scientific papers in DOCX format using python-docx.
#
# Usage: Copy the relevant functions into your paper generation script.
# Dependencies: python-docx

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# =============================================================================
# COLOR PALETTE (Colorblind-safe, from ResNet paper conventions)
# =============================================================================

COLORS = {
    'red': RGBColor(0xE4, 0x1A, 0x1C),
    'green': RGBColor(0x4D, 0xAF, 0x4A),
    'blue': RGBColor(0x37, 0x7E, 0xB8),
    'purple': RGBColor(0x98, 0x4E, 0xA3),
    'orange': RGBColor(0xFF, 0x7F, 0x00),
    'dark': RGBColor(0x33, 0x33, 0x33),
    'gray': RGBColor(0x99, 0x99, 0x99),
}

# =============================================================================
# TYPOGRAPHY HELPERS
# =============================================================================

def set_run_font(run, font_name='Times New Roman', size=10, bold=False,
                 italic=False, color=None):
    """Apply font formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Set East Asian font for CJK compatibility
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_with_style(doc, text, font_name='Times New Roman', size=10,
                             bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, first_line_indent=None):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold, italic=italic)
    return p


# =============================================================================
# HEADING HELPERS
# =============================================================================

def add_section_heading(doc, text, level=1):
    """Add a section heading (1=section, 2=subsection)."""
    sizes = {1: 12, 2: 10}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 10), bold=True)
    return p


# =============================================================================
# FIGURE EMBEDDING (CRITICAL: must be called after text reference)
# =============================================================================

def add_figure(doc, image_path, caption, width=Inches(5.5),
               font_name='Times New Roman', check_exists=True):
    """
    Insert a figure with caption BELOW it.

    CRITICAL: This function must be called immediately after the paragraph
    that references the figure in the text (e.g., "... as shown in Fig. 2 ...").

    Args:
        doc: Document object
        image_path: Absolute or relative path to the image file
        caption: Figure caption text (will be prefixed with "Figure N. ")
        width: Image width (Inches). Use Inches(3.25) for single-column,
               Inches(6.75) for double-column full width.
        font_name: Font for caption
        check_exists: If True, raises FileNotFoundError if image missing

    Returns:
        True if image was inserted, False otherwise

    Raises:
        FileNotFoundError: If check_exists=True and image_path does not exist
    """
    if check_exists and not os.path.exists(image_path):
        raise FileNotFoundError(
            f"Figure image not found: {image_path}\n"
            f"Ensure the image is generated in Phase 2 before embedding in Phase 5."
        )

    # Insert image centered
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(3)
    run = p_img.add_run()
    run.add_picture(image_path, width=width)

    # Caption BELOW the figure (bold, centered)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after = Pt(12)
    run = p_cap.add_run(caption)
    set_run_font(run, font_name=font_name, size=9, bold=True)

    return True


def add_figure_with_fallback(doc, image_path, caption, width=Inches(5.5),
                              fallback_text="[Figure placeholder — image not found]",
                              font_name='Times New Roman'):
    """
    Insert a figure with a text fallback if the image file is missing.
    Use this when image generation might fail but the document should still build.
    """
    if os.path.exists(image_path):
        return add_figure(doc, image_path, caption, width=width, font_name=font_name)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fallback_text)
        set_run_font(run, font_name=font_name, size=9, italic=True, color=COLORS['gray'])
        # Still add caption
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cap.add_run(caption)
        set_run_font(run, font_name=font_name, size=9, bold=True)
        return False


# =============================================================================
# TABLE CREATION
# =============================================================================

def add_table_with_caption(doc, headers, rows, caption,
                           font_name='Times New Roman', bold_best_row=None,
                           best_color=COLORS['blue']):
    """
    Add a table with caption ABOVE it.

    Args:
        doc: Document object
        headers: List of header strings
        rows: List of row lists (each row is a list of cell strings)
        caption: Table caption text (will be prefixed with "Table N. ")
        font_name: Font for table text
        bold_best_row: Index of the row to highlight (typically "Ours" row)
        best_color: Color for the best row text

    Returns:
        The created table object
    """
    # Caption ABOVE the table
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    run = p_cap.add_run(caption)
    set_run_font(run, font_name=font_name, size=9, bold=True)

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, font_name=font_name, size=9, bold=True)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    if bold_best_row is not None and row_idx == bold_best_row:
                        set_run_font(run, font_name=font_name, size=9,
                                     bold=True, color=best_color)
                    else:
                        set_run_font(run, font_name=font_name, size=9)

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    return table


# =============================================================================
# EQUATION PLACEHOLDER
# =============================================================================

def add_equation(doc, latex_text, eq_number, font_name='Cambria Math', size=10):
    """
    Add a numbered equation placeholder.

    Note: python-docx does not support native LaTeX rendering.
    For production papers, use one of:
    1. Pre-render equations as SVG/PNG images and insert via add_picture()
    2. Use docx2pdf + matplotlib for equation rendering
    3. Switch to LaTeX output for camera-ready submission

    This function creates a formatted placeholder that can be replaced
    with actual rendered equations.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Left part: equation text
    run_eq = p.add_run(latex_text)
    set_run_font(run_eq, font_name=font_name, size=size, italic=True)

    # Right part: equation number
    # Use tab to push number to right
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run_tab = p.add_run("\t")
    run_num = p.add_run(f"({eq_number})")
    set_run_font(run_num, font_name=font_name, size=size)

    return p


# =============================================================================
# HEADER AND FOOTER
# =============================================================================

def set_header_footer(doc, title_text, font_name='Times New Roman'):
    """
    Set page header (paper title) and footer (page numbers).

    Args:
        doc: Document object
        title_text: Paper title for header
        font_name: Font for header/footer text
    """
    section = doc.sections[0]

    # Header: paper title (left aligned)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(title_text)
    set_run_font(run, font_name=font_name, size=9, italic=True)

    # Footer: page number (centered)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, font_name=font_name, size=9)


# =============================================================================
# DOCUMENT SETUP
# =============================================================================

def setup_document(title, authors, abstract_text, venue='CVPR',
                   font_name='Times New Roman'):
    """
    Create and initialize a new scientific paper document.

    Returns:
        doc: Initialized Document object
    """
    doc = Document()

    # Page setup: US Letter (8.5" x 11")
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    run = p_title.add_run(title)
    set_run_font(run, font_name=font_name, size=17, bold=True)

    # Authors
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_authors.paragraph_format.space_after = Pt(6)
    run = p_authors.add_run(authors)
    set_run_font(run, font_name=font_name, size=11)

    # Abstract heading
    p_abs_h = doc.add_paragraph()
    p_abs_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_abs_h.paragraph_format.space_before = Pt(12)
    p_abs_h.paragraph_format.space_after = Pt(6)
    run = p_abs_h.add_run("Abstract")
    set_run_font(run, font_name=font_name, size=12, bold=True)

    # Abstract text (single paragraph, justified)
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.space_after = Pt(12)
    run = p_abs.add_run(abstract_text)
    set_run_font(run, font_name=font_name, size=10, italic=False)

    # Set header/footer
    set_header_footer(doc, title, font_name=font_name)

    return doc


# =============================================================================
# COLUMN LAYOUT NOTE
# =============================================================================

def set_double_column_note():
    """
    IMPORTANT NOTE about double-column layout:

    python-docx does NOT natively support double-column (multi-column) layout.
    The DOCX format supports multi-column sections, but python-docx's API
    does not expose this functionality.

    For double-column conference papers (CVPR, ICCV, ECCV, ICLR, ACL), use:

    1. **LaTeX output** (recommended for camera-ready):
       Use the venue's official LaTeX template (e.g., cvpr.sty).
       This is the only way to get true double-column layout with proper
       figure/table placement.

    2. **DOCX workaround** (for collaborative drafting only):
       Generate single-column DOCX with figures at Inches(3.25) width.
       The layout will be single-column, but content is correct.
       Convert to PDF using the venue's Word template if available.

    3. **Post-processing**:
       Generate the DOCX, then manually apply double-column layout in Word,
       or use pandoc to convert to LaTeX:
       pandoc paper.docx -o paper.tex --bibliography=refs.bib

    Therefore:
    - For **drafting/collaboration**: Use DOCX (single column, wider figures)
    - For **camera-ready submission**: Use LaTeX (true double column)
    """
    pass


# =============================================================================
# EXAMPLE USAGE (copy into your generate_paper.py)
# =============================================================================

EXAMPLE_USAGE = '''
# --- Example: Complete paper generation workflow ---

from code_templates import (
    setup_document, add_section_heading, add_paragraph_with_style,
    add_figure, add_table_with_caption, add_equation
)
from docx.shared import Inches

# Phase 1: Plan
venue = 'CVPR'
output_path = 'my_paper.docx'
assets_dir = './paper_assets'

# Phase 2: Generate charts (use matplotlib, save to assets_dir)
# ... generate fig1.png, fig2.png, etc. ...

# Phase 3-5: Draft + Polish + Convert
doc = setup_document(
    title="My Research Paper Title",
    authors="Author One, Author Two",
    abstract_text="This paper presents...",
    venue=venue
)

# 1. Introduction
add_section_heading(doc, "1. Introduction", level=1)
add_paragraph_with_style(doc, "Deep learning has...")
add_paragraph_with_style(doc, "However, a critical gap remains...")
add_paragraph_with_style(doc, "As shown in Fig. 1, our approach...")

# CRITICAL: Embed Figure 1 immediately after its text reference!
add_figure(doc,
    image_path=f"{assets_dir}/fig1_architecture.png",
    caption="Figure 1. System architecture overview. ...",
    width=Inches(5.5))

# Continue with more paragraphs...
add_paragraph_with_style(doc, "We validate our hypothesis...")

# 2. Related Work
add_section_heading(doc, "2. Related Work", level=1)
# ...

# 3. Method
add_section_heading(doc, "3. Method", level=1)
add_section_heading(doc, "3.1 Core Formulation", level=2)
add_equation(doc, "y = f(x) + epsilon", eq_number=1)
# ...

# 4. Experiments
add_section_heading(doc, "4. Experiments", level=1)
add_section_heading(doc, "4.1 Main Benchmark", level=2)

# Add table with caption ABOVE
tbl = add_table_with_caption(
    doc,
    headers=["Method", "Accuracy", "F1"],
    rows=[["Baseline", "85.2", "84.1"], ["Ours", "89.4", "88.7"]],
    caption="Table 1. Main results on XYZ dataset.",
    bold_best_row=1  # Highlight "Ours" row
)

# Embed figure after table reference
add_figure(doc,
    image_path=f"{assets_dir}/fig3_results.png",
    caption="Figure 3. Performance comparison across tasks...",
    width=Inches(5.5))

# 5. Conclusion
add_section_heading(doc, "5. Conclusion", level=1)
# ...

# Save
doc.save(output_path)
print(f"Paper saved to {output_path}")
'''
